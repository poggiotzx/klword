# -*- coding: utf-8 -*-
"""Word 文档生成 API。

本模块基于 ``docxtpl`` 与 ``python-docx`` 提供统一的报告生成接口：

1. ``docxtpl`` 负责模板占位渲染。
2. ``python-docx`` 负责段落、表格、图片、页眉页脚及域的补充写入。
3. 通过 :class:`DocContainer` 提供链式子文档拼装能力。
"""

from __future__ import annotations

import os
import unicodedata
from pathlib import Path
from typing import Any, Optional, Sequence, Union

from docx import Document
from docx.enum.table import (
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_ROW_HEIGHT_RULE,
    WD_TABLE_ALIGNMENT,
)
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docxtpl import DocxTemplate
from PIL import Image

from .word_styles import (
    BODY_STYLE,
    CAPTION_STYLE,
    FOOTER_STYLE,
    HEADER_STYLE,
    IMAGE_STYLE,
    TABLE_BODY_STYLE,
    TABLE_HEADER_STYLE,
    CellStyle,
    TableStyle,
    TextStyle,
    make_cell_style,
    make_table_style,
    make_text_style,
)

TextValue = Union[str, int, float, None]
PartConfig = dict[str, Any]
TableConfig = dict[str, Any]


class DocContainer:
    """子文档容器。"""

    def __init__(self, api: "WordAPI", subdoc: Any) -> None:
        self.api = api
        self.subdoc = subdoc

    def add_title(self, text: str, style: Optional[TextStyle] = None) -> "DocContainer":
        self.api.add_paragraph(
            self.subdoc,
            text,
            style or make_text_style(style_name="KL主标题"),
        )
        return self

    def add_heading(
        self,
        text: str,
        level: int = 1,
        style: Optional[TextStyle] = None,
    ) -> "DocContainer":
        if style is None:
            style = {
                1: make_text_style(style_name="KL一级标题"),
                2: make_text_style(style_name="KL二级标题"),
                3: make_text_style(style_name="KL其他标题"),
            }.get(level, BODY_STYLE)
        self.api.add_paragraph(self.subdoc, text, style)
        return self

    def add_paragraph(
        self,
        text: TextValue,
        style: Optional[TextStyle] = None,
    ) -> "DocContainer":
        self.api.add_paragraph(self.subdoc, text, style or BODY_STYLE)
        return self

    def add_image(
        self,
        image_path: str,
        width_cm: Optional[float] = None,
        height_cm: Optional[float] = None,
        align: str = "center",
        style: Optional[TextStyle] = None,
    ) -> "DocContainer":
        self.api.add_image_block(
            self.subdoc,
            image_path=image_path,
            width_cm=width_cm,
            height_cm=height_cm,
            align=align,
            style=style or IMAGE_STYLE,
        )
        return self

    def add_table(
        self,
        data: list[list[TextValue]],
        header_style: Optional[CellStyle] = None,
        body_style: Optional[CellStyle] = None,
        table_style: Optional[TableStyle] = None,
    ) -> "DocContainer":
        self.api.insert_table(
            container=self.subdoc,
            data=data,
            header_style=header_style,
            body_style=body_style,
            table_style=table_style,
        )
        return self

    def add_table_by_config(self, table_config: TableConfig) -> "DocContainer":
        self.api.insert_table_by_config(self.subdoc, table_config)
        return self

    def add_page_break(self) -> "DocContainer":
        self.subdoc.add_page_break()
        return self

    def add_field_paragraph(
        self,
        parts: list[PartConfig],
        style: Optional[TextStyle] = None,
    ) -> Any:
        return self.api.add_field_paragraph(self.subdoc, parts, style or BODY_STYLE)

    def add_page_footer(self, style: Optional[TextStyle] = None) -> Any:
        return self.api.add_page_footer(self.subdoc, style or FOOTER_STYLE)

    def add_figure_caption_auto(
        self,
        title: str,
        style: Optional[TextStyle] = None,
    ) -> Any:
        return self.api.add_figure_caption_auto(
            self.subdoc,
            title,
            style or CAPTION_STYLE,
        )

    def add_table_caption_auto(
        self,
        title: str,
        style: Optional[TextStyle] = None,
    ) -> Any:
        return self.api.add_table_caption_auto(
            self.subdoc,
            title,
            style or CAPTION_STYLE,
        )


class WordAPI:
    """Word 报告生成主入口。"""

    IMAGE_EXTENSIONS = {
        ".png",
        ".jpg",
        ".jpeg",
        ".bmp",
        ".gif",
        ".webp",
        ".tif",
        ".tiff",
    }

    def __init__(self, template_path: str) -> None:
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"模板不存在: {template_path}")
        self.template_path = str(template_path)
        self.doc = DocxTemplate(self.template_path)

    def new_container(self) -> DocContainer:
        return DocContainer(self, self.doc.new_subdoc())

    @staticmethod
    def _check_color(color: Optional[str]) -> Optional[str]:
        if color is None:
            return None
        value = color.strip().replace("#", "").upper()
        if len(value) != 6:
            raise ValueError(
                "颜色值必须为 6 位十六进制，例如 FF0000，"
                f"当前为: {color}"
            )
        return value

    @staticmethod
    def _get_paragraph_alignment(align: Optional[str]) -> WD_PARAGRAPH_ALIGNMENT:
        value = (align or "left").lower()
        if value == "center":
            return WD_PARAGRAPH_ALIGNMENT.CENTER
        if value == "right":
            return WD_PARAGRAPH_ALIGNMENT.RIGHT
        return WD_PARAGRAPH_ALIGNMENT.LEFT

    @staticmethod
    def _get_vertical_alignment(
        vertical_align: Optional[str],
    ) -> WD_CELL_VERTICAL_ALIGNMENT:
        value = (vertical_align or "center").lower()
        if value == "top":
            return WD_CELL_VERTICAL_ALIGNMENT.TOP
        if value == "bottom":
            return WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
        return WD_CELL_VERTICAL_ALIGNMENT.CENTER

    @staticmethod
    def _get_table_alignment(align: Optional[str]) -> WD_TABLE_ALIGNMENT:
        value = (align or "center").lower()
        if value == "left":
            return WD_TABLE_ALIGNMENT.LEFT
        if value == "right":
            return WD_TABLE_ALIGNMENT.RIGHT
        return WD_TABLE_ALIGNMENT.CENTER

    @staticmethod
    def _set_run_font(
        run: Any,
        font_name: Optional[str] = None,
        font_size: Optional[float] = None,
        bold: Optional[bool] = None,
        italic: Optional[bool] = None,
        font_color: Optional[str] = None,
    ) -> None:
        font_color = WordAPI._check_color(font_color)

        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic
        if font_name:
            run.font.name = font_name
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.rFonts
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            rfonts.set(qn("w:eastAsia"), font_name)
            rfonts.set(qn("w:ascii"), font_name)
            rfonts.set(qn("w:hAnsi"), font_name)
        if font_size is not None:
            run.font.size = Pt(font_size)
        if font_color:
            run.font.color.rgb = RGBColor.from_string(font_color)

    @staticmethod
    def _apply_paragraph_style(paragraph: Any, style_name: Optional[str]) -> None:
        if not style_name:
            return
        try:
            paragraph.style = style_name
        except Exception:
            return

    @staticmethod
    def _apply_paragraph_direct_format(
        paragraph: Any,
        style: Optional[TextStyle],
    ) -> None:
        if style is None:
            return

        if style.align is not None:
            paragraph.alignment = WordAPI._get_paragraph_alignment(style.align)

        paragraph_format = paragraph.paragraph_format
        if style.line_spacing is not None:
            paragraph_format.line_spacing = style.line_spacing
            paragraph_format.line_spacing_rule = (
                WD_LINE_SPACING.SINGLE
                if style.line_spacing == 1
                else WD_LINE_SPACING.MULTIPLE
            )
        if style.space_before_pt is not None:
            paragraph_format.space_before = Pt(style.space_before_pt)
        if style.space_after_pt is not None:
            paragraph_format.space_after = Pt(style.space_after_pt)
        if style.first_line_indent_chars is not None:
            base_pt = style.font_size or 12.0
            paragraph_format.first_line_indent = Pt(
                base_pt * style.first_line_indent_chars
            )

    @staticmethod
    def _set_cell_background(cell: Any, fill: Optional[str]) -> None:
        fill = WordAPI._check_color(fill) if fill else None
        if not fill:
            return
        tc_pr = cell._tc.get_or_add_tcPr()
        shading = tc_pr.find(qn("w:shd"))
        if shading is None:
            shading = OxmlElement("w:shd")
            tc_pr.append(shading)
        shading.set(qn("w:fill"), fill)
        shading.set(qn("w:val"), "clear")

    @staticmethod
    def _set_cell_margins(
        cell: Any,
        top: int = 80,
        start: int = 80,
        bottom: int = 80,
        end: int = 80,
    ) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_margin = tc_pr.find(qn("w:tcMar"))
        if tc_margin is None:
            tc_margin = OxmlElement("w:tcMar")
            tc_pr.append(tc_margin)
        for key, value in {
            "top": top,
            "start": start,
            "bottom": bottom,
            "end": end,
        }.items():
            node = tc_margin.find(qn(f"w:{key}"))
            if node is None:
                node = OxmlElement(f"w:{key}")
                tc_margin.append(node)
            node.set(qn("w:w"), str(value))
            node.set(qn("w:type"), "dxa")

    @staticmethod
    def _cm_to_dxa(cm_value: float) -> int:
        return int(cm_value / 2.54 * 1440)

    @staticmethod
    def _set_cell_width(cell: Any, width_cm: float) -> None:
        width_dxa = WordAPI._cm_to_dxa(width_cm)
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_width = tc_pr.find(qn("w:tcW"))
        if tc_width is None:
            tc_width = OxmlElement("w:tcW")
            tc_pr.append(tc_width)
        tc_width.set(qn("w:w"), str(width_dxa))
        tc_width.set(qn("w:type"), "dxa")

    @staticmethod
    def _set_table_fixed_layout(table: Any) -> None:
        table_pr = table._tbl.tblPr
        table_layout = table_pr.find(qn("w:tblLayout"))
        if table_layout is None:
            table_layout = OxmlElement("w:tblLayout")
            table_pr.append(table_layout)
        table_layout.set(qn("w:type"), "fixed")

    @staticmethod
    def _set_table_grid_widths(table: Any, col_widths_cm: Sequence[float]) -> None:
        table_grid = table._tbl.tblGrid
        if table_grid is None:
            table_grid = OxmlElement("w:tblGrid")
            table._tbl.insert(1, table_grid)
        else:
            for child in list(table_grid):
                table_grid.remove(child)

        for width_cm in col_widths_cm:
            grid_col = OxmlElement("w:gridCol")
            grid_col.set(qn("w:w"), str(WordAPI._cm_to_dxa(float(width_cm))))
            table_grid.append(grid_col)

    @staticmethod
    def _clear_paragraph(paragraph: Any) -> None:
        paragraph_element = paragraph._element
        for child in list(paragraph_element):
            paragraph_element.remove(child)

    @classmethod
    def _is_image_file(cls, path_str: str) -> bool:
        return Path(path_str).suffix.lower() in cls.IMAGE_EXTENSIONS

    @staticmethod
    def _safe_image_size(path_str: str) -> tuple[int, int]:
        with Image.open(path_str) as image:
            return image.size

    def add_empty_paragraph(
        self,
        container: Any,
        style: Optional[TextStyle] = None,
    ) -> Any:
        paragraph = container.add_paragraph()
        self._apply_paragraph_style(paragraph, getattr(style, "style_name", None))
        self._apply_paragraph_direct_format(paragraph, style)
        return paragraph

    def add_text_run(
        self,
        paragraph: Any,
        text: TextValue,
        style: Optional[TextStyle] = None,
    ) -> Any:
        run = paragraph.add_run("" if text is None else str(text))
        if style is not None:
            self._set_run_font(
                run,
                font_name=style.font_name,
                font_size=style.font_size,
                bold=style.bold,
                italic=style.italic,
                font_color=style.font_color,
            )
        return run

    def add_paragraph(
        self,
        container: Any,
        text: TextValue = "",
        style: Optional[TextStyle] = None,
    ) -> Any:
        paragraph = self.add_empty_paragraph(container, style)
        self.add_text_run(paragraph, text, style)
        return paragraph

    def add_image_block(
        self,
        container: Any,
        image_path: str,
        width_cm: Optional[float] = None,
        height_cm: Optional[float] = None,
        align: str = "center",
        style: Optional[TextStyle] = None,
    ) -> Any:
        if not os.path.exists(image_path):
            raise FileNotFoundError(f"图片不存在: {image_path}")

        paragraph = self.add_empty_paragraph(container, style or IMAGE_STYLE)
        paragraph.alignment = self._get_paragraph_alignment(align)
        run = paragraph.add_run()

        if width_cm and height_cm:
            run.add_picture(image_path, width=Cm(width_cm), height=Cm(height_cm))
        elif width_cm:
            run.add_picture(image_path, width=Cm(width_cm))
        elif height_cm:
            run.add_picture(image_path, height=Cm(height_cm))
        else:
            run.add_picture(image_path)

        return paragraph

    def _apply_cell_style(
        self,
        cell: Any,
        style: Optional[CellStyle],
    ) -> None:
        style = style or make_cell_style()
        cell.vertical_alignment = self._get_vertical_alignment(style.vertical_align)
        self._set_cell_margins(cell)

        if style.bg_color:
            self._set_cell_background(cell, style.bg_color)

        paragraphs = cell.paragraphs or [cell.add_paragraph()]
        for index, paragraph in enumerate(paragraphs):
            if index == 0:
                self._clear_paragraph(paragraph)
            self._apply_paragraph_style(paragraph, style.paragraph_style_name)
            if style.align is not None:
                paragraph.alignment = self._get_paragraph_alignment(style.align)
            if style.line_spacing is not None:
                paragraph_format = paragraph.paragraph_format
                paragraph_format.line_spacing = style.line_spacing
                paragraph_format.line_spacing_rule = (
                    WD_LINE_SPACING.SINGLE
                    if style.line_spacing == 1
                    else WD_LINE_SPACING.MULTIPLE
                )

    def _fill_cell_value(
        self,
        cell: Any,
        value: TextValue,
        style: Optional[CellStyle],
        table_style: TableStyle,
    ) -> None:
        self._apply_cell_style(cell, style)
        paragraph = cell.paragraphs[0]

        if isinstance(value, str) and os.path.exists(value) and self._is_image_file(value):
            run = paragraph.add_run()
            if table_style.auto_fit_image:
                image_width_px, image_height_px = self._safe_image_size(value)
                width_cm = table_style.image_width_cm
                height_cm = (
                    width_cm * image_height_px / image_width_px
                    if image_width_px
                    else width_cm
                )
                run.add_picture(value, width=Cm(width_cm), height=Cm(height_cm))
            else:
                run.add_picture(value, width=Cm(table_style.image_width_cm))
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            return

        run = paragraph.add_run("" if value is None else str(value))
        if style is not None:
            self._set_run_font(
                run,
                font_name=style.font_name,
                font_size=style.font_size,
                bold=style.bold,
                italic=style.italic,
                font_color=style.font_color,
            )

    def _set_table_borders(
        self,
        table: Any,
        border_color: str = "000000",
        border_size: str = "8",
    ) -> None:
        table_properties = table._tbl.tblPr
        borders = table_properties.find(qn("w:tblBorders"))
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            table_properties.append(borders)

        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            node = borders.find(qn(f"w:{edge}"))
            if node is None:
                node = OxmlElement(f"w:{edge}")
                borders.append(node)
            node.set(qn("w:val"), "single")
            node.set(qn("w:sz"), str(border_size))
            node.set(qn("w:space"), "0")
            node.set(qn("w:color"), border_color)

    def insert_table(
        self,
        container: Any,
        data: list[list[TextValue]],
        header_style: Optional[CellStyle] = None,
        body_style: Optional[CellStyle] = None,
        table_style: Optional[TableStyle] = None,
    ) -> Any:
        if not data:
            raise ValueError("表格数据不能为空")

        col_count = max(len(row) for row in data)
        row_count = len(data)
        table_style = table_style or make_table_style()
        header_style = header_style or make_cell_style(
            paragraph_style_name="KL表格表头",
            align="center",
        )
        body_style = body_style or make_cell_style(
            paragraph_style_name="KL表格文字",
            align="left",
        )

        table = container.add_table(rows=row_count, cols=col_count)
        table.alignment = self._get_table_alignment(table_style.align)
        table.autofit = False
        self._set_table_fixed_layout(table)
        self._set_table_borders(table, table_style.border_color, table_style.border_size)

        if table_style.col_widths_cm:
            effective_widths = [
                float(width_cm)
                for width_cm in table_style.col_widths_cm[:col_count]
            ]
            if len(effective_widths) < col_count:
                effective_widths.extend(
                    [float(effective_widths[-1])] * (col_count - len(effective_widths))
                )

            self._set_table_grid_widths(table, effective_widths)

            for col_idx, width_cm in enumerate(effective_widths):
                for row in table.rows:
                    self._set_cell_width(row.cells[col_idx], width_cm)

        if table_style.row_heights_cm:
            for row_idx, height_cm in enumerate(table_style.row_heights_cm):
                if row_idx >= row_count or height_cm is None:
                    continue
                row = table.rows[row_idx]
                row.height = Cm(float(height_cm))
                row.height_rule = (
                    WD_ROW_HEIGHT_RULE.EXACTLY
                    if table_style.exact_row_height
                    else WD_ROW_HEIGHT_RULE.AT_LEAST
                )

        for row_index, row_data in enumerate(data):
            for col_index in range(col_count):
                value = row_data[col_index] if col_index < len(row_data) else ""
                style = (
                    header_style
                    if row_index < table_style.header_rows
                    else body_style
                )
                self._fill_cell_value(
                    table.cell(row_index, col_index),
                    value,
                    style,
                    table_style,
                )

        return table

    def insert_table_by_config(self, container: Any, table_config: TableConfig) -> Any:
        default_data = [
            ["序号", "内容"],
            ["未检测到数据填充", "未检测到数据填充"],
        ]

        max_table_width_cm = 14.0
        default_row_height_cm = 0.6
        default_col_width_cm = 4.0
        min_col_width_cm = 1.2

        data = table_config.get("data") or default_data
        data = self._normalize_table_data(data) or default_data

        row_count = len(data)
        input_row_heights = table_config.get("row_heights_cm") or []
        input_col_widths = table_config.get("col_widths_cm") or []

        row_heights_cm = [
            input_row_heights[i] if i < len(input_row_heights) else default_row_height_cm
            for i in range(row_count)
        ]

        col_widths_cm = self._build_col_widths(
            data=data,
            input_col_widths=input_col_widths,
            max_table_width_cm=max_table_width_cm,
            default_col_width_cm=default_col_width_cm,
            min_col_width_cm=min_col_width_cm,
        )

        style_config = table_config.get("style") or {}

        header_style = style_config.get("header", TABLE_HEADER_STYLE)
        body_style = style_config.get("body", TABLE_BODY_STYLE)
        table_style = style_config.get(
            "table",
            make_table_style(
                col_widths_cm=col_widths_cm,
                row_heights_cm=row_heights_cm,
            ),
        )

        return self.insert_table(
            container=container,
            data=data,
            header_style=header_style,
            body_style=body_style,
            table_style=table_style,
        )

    def _normalize_table_data(self, data: list[list[Any]]) -> list[list[str]]:
        if not data:
            return []

        col_count = max(len(row) for row in data)
        normalized = []
        for row in data:
            new_row = ["" if cell is None else str(cell) for cell in row]
            if len(new_row) < col_count:
                new_row.extend([""] * (col_count - len(new_row)))
            normalized.append(new_row)
        return normalized

    def _build_col_widths(
        self,
        data: list[list[str]],
        input_col_widths: list[float],
        max_table_width_cm: float = 14.0,
        default_col_width_cm: float = 4.0,
        min_col_width_cm: float = 1.2,
    ) -> list[float]:
        if not data or not data[0]:
            return []

        col_count = len(data[0])

        col_weights = []
        for col_idx in range(col_count):
            max_len = 1
            for row in data:
                cell_text = row[col_idx] if col_idx < len(row) else ""
                text_len = self._get_display_length(cell_text)
                max_len = max(max_len, text_len)
            col_weights.append(max_len)

        total_weight = sum(col_weights) or col_count
        auto_widths_all = [
            max(min_col_width_cm, max_table_width_cm * weight / total_weight)
            for weight in col_weights
        ]
        auto_widths_all = self._scale_widths_to_max(auto_widths_all, max_table_width_cm)

        specified: list[Optional[float]] = []
        unspecified_indexes = []
        for i in range(col_count):
            if i < len(input_col_widths) and input_col_widths[i] is not None and input_col_widths[i] > 0:
                specified.append(float(input_col_widths[i]))
            else:
                specified.append(None)
                unspecified_indexes.append(i)

        if len(unspecified_indexes) == col_count:
            return auto_widths_all

        if not unspecified_indexes:
            final_widths = [w if w is not None else default_col_width_cm for w in specified]
            return self._scale_widths_to_max(final_widths, max_table_width_cm)

        specified_sum = sum(w for w in specified if w is not None)

        if specified_sum >= max_table_width_cm:
            desired = [
                specified[i] if specified[i] is not None else auto_widths_all[i]
                for i in range(col_count)
            ]
            return self._scale_widths_to_max(desired, max_table_width_cm)

        remaining_width = max_table_width_cm - specified_sum
        unspecified_weights = [col_weights[i] for i in unspecified_indexes]
        unspecified_total_weight = sum(unspecified_weights) or len(unspecified_indexes)

        final_widths = []
        for i in range(col_count):
            if specified[i] is not None:
                final_widths.append(specified[i])  # type: ignore[arg-type]
            else:
                weight = col_weights[i]
                width = remaining_width * weight / unspecified_total_weight
                final_widths.append(max(min_col_width_cm, width))

        return self._scale_widths_to_max(final_widths, max_table_width_cm)

    def _scale_widths_to_max(self, widths: list[float], max_total: float) -> list[float]:
        if not widths:
            return widths

        total = sum(widths)
        if total <= 0 or total <= max_total:
            return widths

        ratio = max_total / total
        return [round(w * ratio, 4) for w in widths]

    def _get_display_length(self, text: Any) -> int:
        s = "" if text is None else str(text)
        length = 0
        for ch in s:
            length += 2 if unicodedata.east_asian_width(ch) in ("F", "W", "A") else 1
        return max(length, 1)

    def add_field_run(self, paragraph: Any, field_code: str) -> Any:
        run_begin = paragraph.add_run()
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        run_begin._r.append(fld_char_begin)

        run_instr = paragraph.add_run()
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = field_code
        run_instr._r.append(instr_text)

        run_sep = paragraph.add_run()
        fld_char_sep = OxmlElement("w:fldChar")
        fld_char_sep.set(qn("w:fldCharType"), "separate")
        run_sep._r.append(fld_char_sep)

        placeholder_run = paragraph.add_run("")

        run_end = paragraph.add_run()
        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")
        run_end._r.append(fld_char_end)
        return placeholder_run

    def add_field_paragraph(
        self,
        container: Any,
        parts: list[PartConfig],
        style: Optional[TextStyle] = None,
    ) -> Any:
        paragraph = self.add_empty_paragraph(container, style or BODY_STYLE)
        run_style = style or BODY_STYLE

        for part in parts:
            part_type = part.get("type", "text")
            if part_type == "text":
                self.add_text_run(paragraph, part.get("value", ""), run_style)
            elif part_type == "field":
                self.add_field_run(paragraph, part.get("code", ""))
            else:
                raise ValueError(f"不支持的 part type: {part_type}")

        return paragraph

    def add_page_footer(
        self,
        container: Any,
        style: Optional[TextStyle] = None,
    ) -> Any:
        return self.add_field_paragraph(
            container,
            [
                {"type": "text", "value": "第 "},
                {"type": "field", "code": "PAGE"},
                {"type": "text", "value": " 页 / 共 "},
                {"type": "field", "code": "NUMPAGES"},
                {"type": "text", "value": " 页"},
            ],
            style or FOOTER_STYLE,
        )

    def add_figure_caption_auto(
        self,
        container: Any,
        title: str,
        style: Optional[TextStyle] = None,
    ) -> Any:
        return self.add_field_paragraph(
            container,
            [
                {"type": "text", "value": "图 "},
                {"type": "field", "code": r"SEQ Figure \* ARABIC"},
                {"type": "text", "value": f" {title}"},
            ],
            style or CAPTION_STYLE,
        )

    def add_table_caption_auto(
        self,
        container: Any,
        title: str,
        style: Optional[TextStyle] = None,
    ) -> Any:
        return self.add_field_paragraph(
            container,
            [
                {"type": "text", "value": "表 "},
                {"type": "field", "code": r"SEQ Table \* ARABIC"},
                {"type": "text", "value": f" {title}"},
            ],
            style or CAPTION_STYLE,
        )

    def render(self, context: dict[str, Any], output_path: str) -> str:
        output_path = str(output_path)
        output_dir = os.path.dirname(output_path)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)

        self.doc.render(context)
        self.doc.save(output_path)
        return output_path

    def _write_single_header(
        self,
        section: Any,
        text: str,
        style: Optional[TextStyle],
    ) -> None:
        header = section.header
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        self._clear_paragraph(paragraph)
        self._apply_paragraph_style(paragraph, getattr(style, "style_name", None))
        self._apply_paragraph_direct_format(paragraph, style)
        paragraph.alignment = self._get_paragraph_alignment(
            getattr(style, "align", "center")
        )
        self.add_text_run(paragraph, text, style)

    def _write_single_footer(
        self,
        section: Any,
        style: Optional[TextStyle],
    ) -> None:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        self._clear_paragraph(paragraph)
        self._apply_paragraph_style(paragraph, getattr(style, "style_name", None))
        self._apply_paragraph_direct_format(paragraph, style)
        paragraph.alignment = self._get_paragraph_alignment(
            getattr(style, "align", "center")
        )

        self.add_text_run(paragraph, "第 ", style)
        self.add_field_run(paragraph, "PAGE")
        self.add_text_run(paragraph, " 页 / 共 ", style)
        self.add_field_run(paragraph, "NUMPAGES")
        self.add_text_run(paragraph, " 页", style)

    def write_header_footer(
        self,
        docx_path: str,
        header_text: Optional[str] = None,
        header_style: Optional[TextStyle] = None,
        footer_style: Optional[TextStyle] = None,
    ) -> str:
        document = Document(docx_path)
        for section in document.sections:
            if header_text is not None:
                self._write_single_header(
                    section,
                    header_text,
                    header_style or HEADER_STYLE,
                )
            self._write_single_footer(section, footer_style or FOOTER_STYLE)

        document.save(docx_path)
        return docx_path


__all__ = [
    "TextValue",
    "PartConfig",
    "TableConfig",
    "DocContainer",
    "WordAPI",
]
